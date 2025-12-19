import pandas as pd
from docx import Document

ruta = "C:/Users/lucia/Downloads/jardineria BASES DE DATOS.docx"
doc = Document(ruta)
texto = []
for p in doc.paragraphs:
    if p.text.strip():
        texto.append(p.text)

contenido = "\n".join(texto)

print(contenido)
print(f"\n📏 Palabras: {len(contenido.split())}")

from docx import Document
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import re


# ============================================
# EXTRAER TABLAS DEL DOCX
# ============================================

def extraer_tablas_docx(ruta_docx):
    """
    Extrae información de tablas de un documento .docx
    """
    doc = Document(ruta_docx)
    tablas_info = []
    tabla_actual = None

    print("=" * 60)
    print("EXTRAYENDO INFORMACIÓN DE TABLAS")
    print("=" * 60)

    for para in doc.paragraphs:
        texto = para.text.strip()

        # Detectar inicio de tabla
        if 'Tabla "' in texto and '":' in texto:
            # Si hay una tabla anterior, guardarla
            if tabla_actual is not None:
                tablas_info.append(tabla_actual)
                print(f"   ✅ Guardada tabla '{tabla_actual['nombre']}' con {len(tabla_actual['atributos'])} atributos")

            # Extraer nombre de la tabla
            match = re.search(r'Tabla "([^"]+)"', texto)
            if match:
                nombre_tabla = match.group(1)
                tabla_actual = {
                    'nombre': nombre_tabla,
                    'atributos': [],
                    'primary_keys': [],
                    'foreign_keys': []
                }
                print(f"\n📋 Encontrada tabla: {nombre_tabla}")

        # Capturar información si estamos dentro de una tabla
        elif tabla_actual is not None and texto:
            # Detectar campos
            if texto.startswith('Tiene los siguientes campos:'):
                continue

            # Detectar clave primaria
            if '(clave primaria' in texto.lower():
                nombre_attr = texto.split('(')[0].strip()
                nombre_attr = re.sub(r'[🔑📌•\-\*\s]+', '', nombre_attr, count=1).strip()

                tabla_actual['primary_keys'].append(nombre_attr)
                tabla_actual['atributos'].append({
                    'nombre': nombre_attr,
                    'tipo': 'PK'
                })
                print(f"  🔑 {nombre_attr} (PK)")

            # Detectar clave foránea
            elif '(clave externa' in texto.lower() or '(clave ajena' in texto.lower():
                nombre_attr = texto.split(':')[0].strip()
                nombre_attr = re.sub(r'[🔑📌•\-\*\s]+', '', nombre_attr, count=1).strip()

                # Buscar tabla referenciada
                match_ref = re.search(r'tabla "([^"]+)"', texto.lower())
                tabla_ref = match_ref.group(1) if match_ref else None

                tabla_actual['foreign_keys'].append({
                    'nombre': nombre_attr,
                    'referencia': tabla_ref
                })
                tabla_actual['atributos'].append({
                    'nombre': nombre_attr,
                    'tipo': 'FK',
                    'referencia': tabla_ref
                })
                print(f"  🔗 {nombre_attr} (FK → {tabla_ref})")

            # Otros atributos normales
            elif ':' in texto:
                nombre_attr = texto.split(':')[0].strip()
                nombre_attr = re.sub(r'[🔑📌•\-\*\s]+', '', nombre_attr, count=1).strip()

                if nombre_attr and len(nombre_attr) < 50:
                    tabla_actual['atributos'].append({
                        'nombre': nombre_attr,
                        'tipo': 'NORMAL'
                    })
                    print(f"  📌 {nombre_attr}")

    # Guardar la última tabla
    if tabla_actual is not None:
        tablas_info.append(tabla_actual)
        print(f"   ✅ Guardada tabla '{tabla_actual['nombre']}' con {len(tabla_actual['atributos'])} atributos")

    print("\n" + "=" * 60)
    print(f"✅ Total de tablas extraídas: {len(tablas_info)}")
    print("=" * 60)

    return tablas_info


# ============================================
# GENERAR DIAGRAMA ER CON MATPLOTLIB
# ============================================

def generar_diagrama_matplotlib(tablas_info, nombre_archivo="diagrama_jardineria"):
    """
    Genera un diagrama ER usando Matplotlib (MEJORADO)
    """
    print("\n🎨 Generando diagrama con Matplotlib...")

    # Configurar figura MÁS GRANDE
    fig, ax = plt.subplots(figsize=(24, 16))  # ← Aumentado de (20, 14)
    ax.set_xlim(0, 16)  # ← Aumentado de 12
    ax.set_ylim(0, 12)  # ← Aumentado de 10
    ax.axis('off')

    # Calcular disposición de tablas (MÁS ESPACIO)
    num_tablas = len(tablas_info)
    cols = 4  # 4 columnas
    rows = (num_tablas + cols - 1) // cols

    posiciones = {}
    relaciones = []

    idx = 0

    for i in range(rows):
        for j in range(cols):
            if idx >= num_tablas:
                break

            tabla = tablas_info[idx]

            # Calcular posición (MÁS ESPACIO ENTRE TABLAS)
            x = 0.8 + j * 4  # ← Cambiado de 3 a 4 (más espacio horizontal)
            y = 11 - i * 3.5  # ← Cambiado de 2.5 a 3.5 (más espacio vertical)

            # Calcular altura y ancho (TABLAS MÁS GRANDES)
            num_attrs = len(tabla['atributos'])
            altura = 0.4 + num_attrs * 0.18  # ← Aumentado de 0.15 a 0.18
            ancho = 3.2  # ← Aumentado de 2.5 a 3.2

            # Dibujar rectángulo de la tabla
            rect = FancyBboxPatch(
                (x, y - altura),
                ancho,
                altura,
                boxstyle="round,pad=0.08",
                edgecolor='#2C3E50',
                facecolor='#ECF0F1',
                linewidth=3  # ← Aumentado de 2.5 a 3
            )
            ax.add_patch(rect)

            # Nombre de la tabla (header) - MÁS ALTO
            header_altura = 0.35  # ← Aumentado de 0.25 a 0.35
            header_rect = FancyBboxPatch(
                (x, y - header_altura),
                ancho,
                header_altura,
                boxstyle="round,pad=0.04",
                edgecolor='#2C3E50',
                facecolor='#3498DB',
                linewidth=2.5
            )
            ax.add_patch(header_rect)

            ax.text(
                x + ancho / 2,
                y - header_altura / 2 + 0.02,
                tabla['nombre'].upper(),
                fontsize=12,  # ← Aumentado de 11 a 13
                fontweight='bold',
                ha='center',
                va='center',
                color='white',
                family='monospace'
            )

            # Atributos (MÁS GRANDES)
            y_attr = y - header_altura - 0.25
            for attr in tabla['atributos']:
                # Determinar estilo según tipo
                if attr['tipo'] == 'PK':
                    prefijo = '[PK] '
                    color = '#E74C3C'
                    weight = 'bold'
                elif attr['tipo'] == 'FK':
                    prefijo = '[FK] '
                    color = '#3498DB'
                    weight = 'bold'

                    # Guardar relación
                    if 'referencia' in attr and attr['referencia']:
                        relaciones.append({
                            'origen': tabla['nombre'],
                            'destino': attr['referencia'],
                            'atributo': attr['nombre']
                        })
                else:
                    prefijo = '     '
                    color = '#2C3E50'
                    weight = 'normal'

                ax.text(
                    x + 0.15,
                    y_attr,
                    f"{prefijo}{attr['nombre']}",
                    fontsize=12,  # ← Aumentado de 8 a 10
                    ha='left',
                    va='center',
                    color=color,
                    weight=weight,
                    family='monospace'
                )
                y_attr -= 0.20  # ← Aumentado

            # Guardar posición central de la tabla
            posiciones[tabla['nombre']] = (x + ancho / 2, y - altura / 2)

            idx += 1

    # Dibujar relaciones (flechas FK)
    for rel in relaciones:
        origen = rel['origen']
        destino = rel['destino']

        if origen in posiciones and destino in posiciones:
            x1, y1 = posiciones[origen]
            x2, y2 = posiciones[destino]

            # Crear flecha
            arrow = FancyArrowPatch(
                (x1, y1),
                (x2, y2),
                arrowstyle='->',
                color='#E74C3C',
                linewidth=2,  # ← Aumentado de 1.5 a 2
                linestyle='--',
                mutation_scale=25,  # ← Aumentado de 20 a 25
                alpha=0.7
            )
            ax.add_patch(arrow)

            # Etiqueta de la relación
            mid_x = (x1 + x2) / 2
            mid_y = (y1 + y2) / 2
            ax.text(
                mid_x,
                mid_y,
                rel['atributo'],
                fontsize=10,  # ← Aumentado de 7 a 9
                ha='center',
                bbox=dict(
                    boxstyle='round,pad=0.4',
                    facecolor='white',
                    edgecolor='#E74C3C',
                    alpha=0.9
                )
            )

    # Título
    plt.suptitle(
        'Diagrama Entidad-Relación - Base de Datos Jardinería',
        fontsize=20,  # ← Aumentado de 18 a 20
        fontweight='bold',
        y=0.98
    )

    # Leyenda (ESQUINA INFERIOR IZQUIERDA)
    legend_elements = [
        mpatches.Patch(facecolor='#E74C3C', edgecolor='black', label='[PK] Clave Primaria'),
        mpatches.Patch(facecolor='#3498DB', edgecolor='black', label='[FK] Clave Foránea'),
        mpatches.Patch(facecolor='#ECF0F1', edgecolor='black', label='Atributo Normal')
    ]
    ax.legend(
        handles=legend_elements,
        loc='lower left',  # ← CAMBIADO de 'upper left' a 'lower left'
        fontsize=11,  # ← Aumentado de 10 a 11
        frameon=True,
        fancybox=True,
        shadow=True
    )

    # Guardar
    plt.tight_layout()
    plt.savefig(f'{nombre_archivo}.png', dpi=300, bbox_inches='tight', facecolor='white')
    print(f"✅ Diagrama guardado como '{nombre_archivo}.png'")

    # También guardar en PDF
    plt.savefig(f'{nombre_archivo}.pdf', dpi=300, bbox_inches='tight', facecolor='white')
    print(f"✅ Diagrama guardado como '{nombre_archivo}.pdf'")

    plt.show()

# ============================================
# EJECUTAR
# ============================================

if __name__ == "__main__":
    # Ruta al archivo DOCX
    ruta_docx = "C:/Users/lucia/Downloads/jardineria BASES DE DATOS.docx"

    # Extraer tablas
    print("🔍 Procesando documento...\n")
    tablas = extraer_tablas_docx(ruta_docx)

    # Verificar que se extrajeron tablas
    if not tablas:
        print("❌ No se encontraron tablas en el documento")
        exit()

    # Mostrar resumen
    print("\n" + "=" * 60)
    print("📊 RESUMEN DETALLADO")
    print("=" * 60)

    for i, tabla in enumerate(tablas, 1):
        print(f"\n{i}. Tabla: {tabla['nombre']}")
        print(f"   - Total atributos: {len(tabla['atributos'])}")
        print(f"   - Claves primarias: {', '.join(tabla['primary_keys']) if tabla['primary_keys'] else 'Ninguna'}")
        print(f"   - Claves foráneas: {len(tabla['foreign_keys'])}")

        if tabla['foreign_keys']:
            for fk in tabla['foreign_keys']:
                print(f"      └─ {fk['nombre']} → {fk['referencia']}")

    # Generar diagrama
    print("\n" + "=" * 60)
    print("🎨 GENERANDO DIAGRAMA ER CON MATPLOTLIB")
    print("=" * 60)

    generar_diagrama_matplotlib(tablas, "diagrama_jardineria_matplotlib")

    print("\n✅ ¡Proceso completado exitosamente!")
    print(f"\n📁 Archivos generados:")
    print(f"  - diagrama_jardineria_matplotlib.png")
    print(f"  - diagrama_jardineria_matplotlib.pdf")

    # Verificación
    print(f"\n🔍 Verificación:")
    print(f"Tablas extraídas: {[t['nombre'] for t in tablas]}")